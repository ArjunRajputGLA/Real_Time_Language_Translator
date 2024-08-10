
import streamlit as st
import time 
from google.cloud import translate_v2 as translate
import fitz  
from docx import Document
import os
import re 
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from io import BytesIO
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph
import base64
from gtts import gTTS
import tempfile

st.set_page_config(
    page_title="Document Translator",
    page_icon='🔊',
)

def get_base64_of_video(video_file):
    with open(video_file, "rb") as f:
        data = f.read()
    return base64.b64encode(data).decode()

def text_to_speech(text, lang='en'):
    tts = gTTS(text=text, lang=lang)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as temp_audio:
        tts.save(temp_audio.name)
        return temp_audio.name

logo_path = "Logo.mp4"
if os.path.exists(logo_path):
    video_base64 = get_base64_of_video(logo_path)
    st.sidebar.markdown(
        f"""
        <style>
        [data-testid="stSidebar"] {{
            padding-top: 0rem;
        }}
        .sidebar-logo {{
            position: relative;
            top: -1rem;  
            width: 100%;
            margin-bottom: -3rem;  
        }}
        .sidebar-logo video {{
            width: 100%;
        }}
        </style>
        <div class="sidebar-logo">
            <video autoplay loop muted playsinline>
                <source src="data:video/mp4;base64,{video_base64}" type="video/mp4">
            </video>
        </div>
        """,
        unsafe_allow_html=True
    )
else:
    st.sidebar.warning("Logo video not found. Please ensure 'Logo.mp4' is in the same directory as the script.") 
st.sidebar.markdown("---")

os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = 'language-translator-430016-11c0f67992a4.json' 
translate_client = translate.Client() 

LANGUAGES = {
    'af': 'Afrikaans', 'sq': 'Albanian', 'am': 'Amharic', 'ar': 'Arabic', 'hy': 'Armenian',
    'az': 'Azerbaijani', 'eu': 'Basque', 'be': 'Belarusian', 'bn': 'Bengali', 'bs': 'Bosnian',
    'bg': 'Bulgarian', 'ca': 'Catalan', 'ceb': 'Cebuano', 'ny': 'Chichewa', 'zh-cn': 'Chinese (Simplified)',
    'zh-tw': 'Chinese (Traditional)', 'co': 'Corsican', 'hr': 'Croatian', 'cs': 'Czech', 'da': 'Danish',
    'nl': 'Dutch', 'en': 'English', 'eo': 'Esperanto', 'et': 'Estonian', 'tl': 'Filipino', 'fi': 'Finnish',
    'fr': 'French', 'fy': 'Frisian', 'gl': 'Galician', 'ka': 'Georgian', 'de': 'German', 'el': 'Greek',
    'gu': 'Gujarati', 'ht': 'Haitian Creole', 'ha': 'Hausa', 'haw': 'Hawaiian', 'iw': 'Hebrew', 'hi': 'Hindi',
    'hmn': 'Hmong', 'hu': 'Hungarian', 'is': 'Icelandic', 'ig': 'Igbo', 'id': 'Indonesian', 'ga': 'Irish',
    'it': 'Italian', 'ja': 'Japanese', 'jw': 'Javanese', 'kn': 'Kannada', 'kk': 'Kazakh', 'km': 'Khmer',
    'ko': 'Korean', 'ku': 'Kurdish (Kurmanji)', 'ky': 'Kyrgyz', 'lo': 'Lao', 'la': 'Latin', 'lv': 'Latvian',
    'lt': 'Lithuanian', 'lb': 'Luxembourgish', 'mk': 'Macedonian', 'mg': 'Malagasy', 'ms': 'Malay',
    'ml': 'Malayalam', 'mt': 'Maltese', 'mi': 'Maori', 'mr': 'Marathi', 'mn': 'Mongolian', 'my': 'Myanmar (Burmese)',
    'ne': 'Nepali', 'no': 'Norwegian', 'or': 'Odia', 'ps': 'Pashto', 'fa': 'Persian', 'pl': 'Polish',
    'pt': 'Portuguese', 'pa': 'Punjabi', 'ro': 'Romanian', 'ru': 'Russian', 'sm': 'Samoan', 'gd': 'Scots Gaelic',
    'sr': 'Serbian', 'st': 'Sesotho', 'sn': 'Shona', 'sd': 'Sindhi', 'si': 'Sinhala', 'sk': 'Slovak',
    'sl': 'Slovenian', 'so': 'Somali', 'es': 'Spanish', 'su': 'Sundanese', 'sw': 'Swahili', 'sv': 'Swedish',
    'tg': 'Tajik', 'ta': 'Tamil', 'te': 'Telugu', 'th': 'Thai', 'tr': 'Turkish', 'uk': 'Ukrainian', 'ur': 'Urdu',
    'ug': 'Uyghur', 'uz': 'Uzbek', 'vi': 'Vietnamese', 'cy': 'Welsh', 'xh': 'Xhosa', 'yi': 'Yiddish',
    'yo': 'Yoruba', 'zu': 'Zulu'
}

def translate_text(text, dest_lang):
    tags = re.findall(r'<[^>]+>', text)
    text_with_placeholders = re.sub(r'<[^>]+>', '{}', text)
    result = translate_client.translate(text_with_placeholders, target_language=dest_lang)
    translated_text = result['translatedText']
    for tag in tags:
        translated_text = translated_text.replace('{}', tag, 1)
    return translated_text

def file_contains_text(file):
    file.seek(0)  # Reset file pointer to the beginning
    if file.type == "application/pdf":
        try:
            document = fitz.open(stream=file.read(), filetype="pdf")
            text = ""
            for page in document:
                text += page.get_text("text")
            return bool(text.strip())
        except Exception:
            return False
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            doc = Document(file)
            text = "\n".join(para.text for para in doc.paragraphs)
            return bool(text.strip())
        except Exception:
            return False
    elif file.type == "text/plain":
        try:
            text = file.getvalue().decode("utf-8")
            return bool(text.strip())
        except Exception:
            return False
    else:
        return False

def process_pdf(file):
    file.seek(0)  # Reset file pointer to the beginning
    document = fitz.open(stream=file.read(), filetype="pdf")
    text = ""
    for page in document:
        text += page.get_text("text")
    return text

def process_docx(file):
    file.seek(0)  # Reset file pointer to the beginning
    doc = Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text, doc

def create_pdf_from_text(text, file_path):
    doc = SimpleDocTemplate(file_path, pagesize=letter)
    styles = getSampleStyleSheet()
    flowables = []
    for para in text.split('\n'):
        if para.strip():
            p = Paragraph(para, styles['Normal'])
            flowables.append(p)
    doc.build(flowables)

def create_docx_from_text(text):
    doc = Document()
    for para in text.split('\n'):
        if para.strip():
            doc.add_paragraph(para)
    return doc

def create_txt_from_text(text, file_path):
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(text)

def main():
    st.markdown("""
        <style>
        .big-font {
            font-size: 60px;
            font-weight: bold;
            color: #1E90FF;
            margin-top: -65px;  
            margin-bottom: 70px;
            margin-left: -170px; 
        }
        .medium-font {
            font-size:25px;
            color: lime;
            font-weight: bold;
        }
        .document-text {
            font-size:25px;
            color: cyan;
            font-weight: bold;
        }
        .stWarning {
            background-color: #fff3cd;
            color: #856404;
            padding: 10px;
            border-radius: 4px;
            border: 1px solid #ffeeba;
        }
        .stButton > button {
            width: 100%;
        }
        </style>
        """, unsafe_allow_html=True)

    st.markdown('<p class="big-font">Document 📃 Translator</p>', unsafe_allow_html=True)

    language_options = {v: k for k, v in LANGUAGES.items()}

    with st.sidebar:
        st.markdown(" ")
        st.markdown(" ")
        st.markdown('<p class="medium-font">Select Target Language</p>', unsafe_allow_html=True)
        dest_lang_name = st.selectbox("", options=sorted(language_options.keys()))
        st.info("⚠️ Please select a target language")
        dest_lang_code = language_options[dest_lang_name]
        st.markdown("---")

    if 'translated_text' not in st.session_state:
        st.session_state.translated_text = None
    if 'file_name' not in st.session_state:
        st.session_state.file_name = None

    uploaded_file = st.file_uploader("Upload a file (PDF, DOCX, or TXT)", type=["pdf", "docx", "txt"])
    if uploaded_file is not None:
        if not file_contains_text(uploaded_file):
            st.warning("The uploaded file appears to be empty or does not contain any text. Please upload a file with text content.")
        else:
            file_type = uploaded_file.name.split('.')[-1]
            st.session_state.file_name = uploaded_file.name.split('.')[0]
            if file_type == 'pdf':
                text = process_pdf(uploaded_file) 
            elif file_type == 'docx':
                text, original_doc = process_docx(uploaded_file)
            elif file_type == 'txt':
                uploaded_file.seek(0)  # Reset file pointer to the beginning
                text = uploaded_file.getvalue().decode("utf-8")
            else:
                st.error("Unsupported file type")
                return

            st.markdown("---")
            st.markdown('<p class="document-text">Original Text :</p>', unsafe_allow_html=True)
            st.text_area("", text, height = 220)  
            
            st.markdown("---")

            if st.button("Translate") or st.session_state.translated_text is not None:
                if st.session_state.translated_text is None:
                    st.session_state.translated_text = translate_text(text, dest_lang_code)
                
                st.markdown("---")
                st.markdown('<p class="document-text">Translated Text :</p>', unsafe_allow_html=True)
                st.text_area("", st.session_state.translated_text, height=200)

                st.markdown("---")

                space0, col1, space1, col2, space2, col3 = st.columns([0.1, 1, 0.33, 1, 0.4, 1])

                with space0:
                    st.write(" ")
                pdf_output = BytesIO()
                create_pdf_from_text(st.session_state.translated_text, pdf_output)
                pdf_output.seek(0)
                with col1:
                    st.download_button(
                        label="Download as PDF",
                        data=pdf_output,
                        file_name=f"{st.session_state.file_name}(translated).pdf",
                        mime="application/pdf"
                    )

                with space1:
                    st.write(" ")
                # DOCX download
                docx_output = BytesIO()
                doc = create_docx_from_text(st.session_state.translated_text)
                doc.save(docx_output)
                docx_output.seek(0)
                with col2:
                    st.download_button(
                        label="Download as DOCX",
                        data=docx_output.getvalue(),
                        file_name=f"{st.session_state.file_name}(translated).docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                with space2:
                    st.write(" ")
                # TXT download
                with col3:
                    st.download_button(
                        label="Download as TXT",
                        data=st.session_state.translated_text.encode('utf-8'),
                        file_name=f"{st.session_state.file_name}(translated).txt",
                        mime="text/plain"
                    )
                    
    else:
        st.info("Please upload a file to translate.")

if __name__ == "__main__":
    main() 